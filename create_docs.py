#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_documentation():
    doc = Document()
    
    # Title Page
    title = doc.add_heading('PixelSqueeze', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Cloud Image Optimization Platform', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('A comprehensive SaaS platform for image optimization with API service for third-party integrations.')
    doc.add_paragraph('Version: 1.0.0')
    doc.add_paragraph('Documentation Date: January 2025')
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc = [
        '1. Executive Summary',
        '2. Project Overview', 
        '3. Architecture',
        '4. Technology Stack',
        '5. Installation & Setup',
        '6. API Documentation',
        '7. Frontend Components',
        '8. Backend Services',
        '9. Database Models',
        '10. Payment Integration',
        '11. Deployment',
        '12. Testing',
        '13. Security',
        '14. Monitoring & Analytics',
        '15. Troubleshooting',
        '16. Development Guidelines'
    ]
    
    for item in toc:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('1. Executive Summary', 1)
    doc.add_paragraph('PixelSqueeze is a comprehensive cloud-based image optimization platform designed to provide professional-grade image compression services through both a web interface and REST API. The platform serves as a SaaS solution for businesses and developers who need to optimize images at scale while maintaining quality and performance.')
    
    doc.add_heading('Key Features', 2)
    features = [
        'Web-based drag-and-drop image optimization interface',
        'REST API for third-party integrations',
        'Real-time image optimization (JPG/PNG)',
        'Before/After size comparison',
        'Bulk download as ZIP',
        'User authentication and usage statistics',
        'Subscription-based pricing plans',
        'Secure payment processing via Stripe',
        'AWS S3 integration for scalable storage',
        'Redis caching for performance optimization'
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # Project Overview
    doc.add_heading('2. Project Overview', 1)
    doc.add_heading('Subscription Plans', 2)
    plans = [
        'Free Tier: 100 images/month, max 2MB each',
        'Starter: 5,000 images - $9/month',
        'Pro: 20,000 images - $29/month',
        'Enterprise: 100,000 images - $99/month'
    ]
    
    for plan in plans:
        doc.add_paragraph(plan, style='List Bullet')
    
    # Technology Stack
    doc.add_heading('3. Technology Stack', 1)
    doc.add_heading('Backend Technologies', 2)
    backend_tech = [
        'Node.js 18+ - JavaScript runtime',
        'Express.js - Web framework',
        'MongoDB - NoSQL database',
        'Mongoose - MongoDB ODM',
        'Redis - In-memory cache',
        'Sharp - High-performance image processing',
        'Imagemin - Image optimization',
        'Stripe - Payment processing',
        'AWS SDK - Cloud services integration'
    ]
    
    for tech in backend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_heading('Frontend Technologies', 2)
    frontend_tech = [
        'Next.js 14 - React framework',
        'React 18 - UI library',
        'TypeScript - Type safety',
        'Tailwind CSS - Utility-first CSS',
        'Framer Motion - Animations',
        'React Query - Data fetching',
        'Zustand - State management'
    ]
    
    for tech in frontend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    # Installation & Setup
    doc.add_heading('4. Installation & Setup', 1)
    doc.add_heading('Prerequisites', 2)
    prereqs = [
        'Node.js 18+ installed',
        'MongoDB instance (local or cloud)',
        'AWS S3 bucket configured',
        'Stripe account for payments',
        'Redis instance (optional)'
    ]
    
    for prereq in prereqs:
        doc.add_paragraph(prereq, style='List Bullet')
    
    doc.add_heading('Quick Installation', 2)
    doc.add_paragraph('1. Clone the repository:')
    doc.add_paragraph('git clone <repository-url>', style='List Number')
    
    doc.add_paragraph('2. Install dependencies:')
    doc.add_paragraph('npm install && cd client && npm install', style='List Number')
    
    doc.add_paragraph('3. Configure environment:')
    doc.add_paragraph('cp .env.example .env', style='List Number')
    
    # API Documentation
    doc.add_heading('5. API Documentation', 1)
    doc.add_heading('Authentication', 2)
    doc.add_paragraph('All API requests require an API key in the Authorization header:')
    doc.add_paragraph('Authorization: Bearer YOUR_API_KEY', style='List Number')
    
    doc.add_heading('Endpoints', 2)
    doc.add_heading('POST /api/optimize', 3)
    doc.add_paragraph('Upload and optimize an image file.')
    
    doc.add_heading('POST /api/optimize-url', 3)
    doc.add_paragraph('Optimize an image from a URL.')
    
    doc.add_heading('GET /api/stats', 3)
    doc.add_paragraph('Get user\'s monthly usage statistics.')
    
    # Frontend Components
    doc.add_heading('6. Frontend Components', 1)
    doc.add_heading('Core Components', 2)
    components = [
        'Layout.tsx - Main application layout',
        'ImageUploader.tsx - Drag-and-drop upload',
        'BeforeAfter.tsx - Image comparison',
        'RecentImages.tsx - User image history',
        'BillingPlans.tsx - Subscription selection',
        'AuthProvider.tsx - Authentication context'
    ]
    
    for component in components:
        doc.add_paragraph(component, style='List Bullet')
    
    # Backend Services
    doc.add_heading('7. Backend Services', 1)
    doc.add_heading('Core Services', 2)
    services = [
        'imageProcessor.js - Core optimization logic',
        'storageService.js - AWS S3 management',
        'stripeService.js - Payment processing',
        'emailService.js - Transactional emails',
        'cleanupJob.js - Automated cleanup',
        'logger.js - Centralized logging'
    ]
    
    for service in services:
        doc.add_paragraph(service, style='List Bullet')
    
    # Database Models
    doc.add_heading('8. Database Models', 1)
    doc.add_heading('User Model', 2)
    user_fields = [
        'email - User email (unique)',
        'password - Hashed password',
        'subscription - Current plan',
        'apiKey - API access key',
        'usage - Monthly statistics'
    ]
    
    for field in user_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    doc.add_heading('Image Model', 2)
    image_fields = [
        'userId - User reference',
        'originalSize - Original file size',
        'optimizedSize - Optimized size',
        'downloadUrl - S3 download link',
        'expiresAt - Link expiration'
    ]
    
    for field in image_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    # Payment Integration
    doc.add_heading('9. Payment Integration', 1)
    doc.add_heading('Stripe Integration', 2)
    doc.add_paragraph('PixelSqueeze uses Stripe for subscription management and payment processing.')
    
    doc.add_heading('Webhook Events', 2)
    webhooks = [
        'customer.subscription.created',
        'customer.subscription.updated',
        'customer.subscription.deleted',
        'invoice.payment_succeeded',
        'invoice.payment_failed'
    ]
    
    for webhook in webhooks:
        doc.add_paragraph(webhook, style='List Bullet')
    
    # Deployment
    doc.add_heading('10. Deployment', 1)
    doc.add_heading('Production Deployment', 2)
    deployment_info = [
        'Backend: Deployed on Render.com',
        'Frontend: Deployed on Vercel',
        'Database: MongoDB Atlas',
        'Storage: AWS S3',
        'Cache: Redis Cloud (optional)'
    ]
    
    for info in deployment_info:
        doc.add_paragraph(info, style='List Bullet')
    
    # Security
    doc.add_heading('11. Security', 1)
    doc.add_heading('Security Measures', 2)
    security_measures = [
        'JWT Authentication',
        'API Key Management',
        'Rate Limiting',
        'Input Validation',
        'HTTPS Enforcement',
        'Environment Variables'
    ]
    
    for measure in security_measures:
        doc.add_paragraph(measure, style='List Bullet')
    
    # Troubleshooting
    doc.add_heading('12. Troubleshooting', 1)
    doc.add_heading('Common Issues', 2)
    
    doc.add_heading('Redis Connection Issues', 3)
    doc.add_paragraph('If you see Redis connection warnings:')
    redis_solutions = [
        'Check if Redis service is running',
        'Verify Redis connection URL',
        'Application will fall back to memory cache'
    ]
    
    for solution in redis_solutions:
        doc.add_paragraph(solution, style='List Bullet')
    
    # Development Guidelines
    doc.add_heading('13. Development Guidelines', 1)
    doc.add_heading('Code Standards', 2)
    standards = [
        'Use TypeScript for type safety',
        'Follow ESLint configuration',
        'Write meaningful commit messages',
        'Include proper error handling',
        'Use consistent naming conventions'
    ]
    
    for standard in standards:
        doc.add_paragraph(standard, style='List Bullet')
    
    # Save document
    doc.save('PixelSqueeze_Project_Documentation.docx')
    print("Documentation created successfully!")

if __name__ == "__main__":
    create_documentation()
