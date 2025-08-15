#!/usr/bin/env python3
"""
PixelSqueeze Project Documentation Generator
Creates a comprehensive DOCX file with all project information
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import os
import json

def create_documentation():
    # Create a new Document
    doc = Document()
    
    # Set up styles
    setup_styles(doc)
    
    # Title Page
    add_title_page(doc)
    
    # Table of Contents
    add_table_of_contents(doc)
    
    # Executive Summary
    add_executive_summary(doc)
    
    # Project Overview
    add_project_overview(doc)
    
    # Architecture
    add_architecture(doc)
    
    # Technology Stack
    add_technology_stack(doc)
    
    # Installation & Setup
    add_installation_setup(doc)
    
    # API Documentation
    add_api_documentation(doc)
    
    # Frontend Components
    add_frontend_components(doc)
    
    # Backend Services
    add_backend_services(doc)
    
    # Database Models
    add_database_models(doc)
    
    # Payment Integration
    add_payment_integration(doc)
    
    # Deployment
    add_deployment(doc)
    
    # Testing
    add_testing(doc)
    
    # Security
    add_security(doc)
    
    # Monitoring & Analytics
    add_monitoring_analytics(doc)
    
    # Troubleshooting
    add_troubleshooting(doc)
    
    # Development Guidelines
    add_development_guidelines(doc)
    
    # Save the document
    doc.save('PixelSqueeze_Project_Documentation.docx')
    print("✅ Documentation created successfully: PixelSqueeze_Project_Documentation.docx")

def setup_styles(doc):
    """Set up custom styles for the document"""
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Heading 1 style
    h1_style = styles.add_style('CustomH1', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = None  # Default color
    
    # Heading 2 style
    h2_style = styles.add_style('CustomH2', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.size = Pt(16)
    h2_style.font.bold = True
    
    # Code style
    code_style = styles.add_style('CodeStyle', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(10)

def add_title_page(doc):
    """Add the title page"""
    # Title
    title = doc.add_paragraph('PixelSqueeze', style='CustomTitle')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Cloud Image Optimization Platform', style='CustomH2')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Description
    doc.add_paragraph()
    description = doc.add_paragraph('A fully web-based SaaS platform for image optimization with API service for third-party integrations.')
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Version and Date
    doc.add_paragraph()
    version = doc.add_paragraph('Version: 1.0.0')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date = doc.add_paragraph('Documentation Date: January 2025')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Page break
    doc.add_page_break()

def add_table_of_contents(doc):
    """Add table of contents"""
    doc.add_heading('Table of Contents', level=1)
    
    toc_items = [
        '1. Executive Summary',
        '2. Project Overview',
        '3. Architecture',
        '4. Technology Stack',
        '5. Installation & Setup',
        '6. API Documentation',
        '7. Frontend Components',
        '8. Backend Services',
        '9. Database Models',
        '10. Payment Integration',
        '11. Deployment',
        '12. Testing',
        '13. Security',
        '14. Monitoring & Analytics',
        '15. Troubleshooting',
        '16. Development Guidelines'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()

def add_executive_summary(doc):
    """Add executive summary section"""
    doc.add_heading('1. Executive Summary', level=1)
    
    doc.add_paragraph('PixelSqueeze is a comprehensive cloud-based image optimization platform designed to provide professional-grade image compression services through both a web interface and REST API. The platform serves as a SaaS solution for businesses and developers who need to optimize images at scale while maintaining quality and performance.')
    
    doc.add_heading('Key Features', level=2)
    features = [
        'Web-based drag-and-drop image optimization interface',
        'REST API for third-party integrations',
        'Real-time image optimization (JPG/PNG)',
        'Before/After size comparison',
        'Bulk download as ZIP',
        'User authentication and usage statistics',
        'Subscription-based pricing plans',
        'Secure payment processing via Stripe',
        'AWS S3 integration for scalable storage',
        'Redis caching for performance optimization'
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('Business Value', level=2)
    doc.add_paragraph('The platform addresses the growing need for efficient image optimization in web applications, mobile apps, and content management systems. By providing both a user-friendly web interface and a robust API, PixelSqueeze serves both end-users and developers, creating multiple revenue streams through subscription plans and API usage.')

def add_project_overview(doc):
    """Add project overview section"""
    doc.add_heading('2. Project Overview', level=1)
    
    doc.add_heading('Project Goals', level=2)
    goals = [
        'Create a scalable image optimization service',
        'Provide both web interface and API access',
        'Implement subscription-based monetization',
        'Ensure high performance and reliability',
        'Support multiple image formats and optimization levels',
        'Integrate with popular cloud services'
    ]
    
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')
    
    doc.add_heading('Target Audience', level=2)
    audiences = [
        'Web developers and agencies',
        'E-commerce businesses',
        'Content creators and bloggers',
        'Mobile app developers',
        'Marketing teams',
        'Individual users needing image optimization'
    ]
    
    for audience in audiences:
        doc.add_paragraph(audience, style='List Bullet')
    
    doc.add_heading('Subscription Plans', level=2)
    plans = [
        'Free Tier: 100 images/month, max 2MB each',
        'Starter: 5,000 images - $9/month',
        'Pro: 20,000 images - $29/month',
        'Enterprise: 100,000 images - $99/month'
    ]
    
    for plan in plans:
        doc.add_paragraph(plan, style='List Bullet')

def add_architecture(doc):
    """Add architecture section"""
    doc.add_heading('3. Architecture', level=1)
    
    doc.add_heading('System Architecture', level=2)
    doc.add_paragraph('PixelSqueeze follows a modern microservices architecture with clear separation of concerns:')
    
    architecture_components = [
        'Frontend: Next.js React application with TypeScript',
        'Backend: Node.js Express server with REST API',
        'Database: MongoDB for data persistence',
        'Cache: Redis for performance optimization',
        'Storage: AWS S3 for file storage',
        'Payments: Stripe for subscription management',
        'Email: Nodemailer for transactional emails',
        'Monitoring: Winston for logging and Sentry for error tracking'
    ]
    
    for component in architecture_components:
        doc.add_paragraph(component, style='List Bullet')
    
    doc.add_heading('Data Flow', level=2)
    doc.add_paragraph('1. User uploads image through web interface or API')
    doc.add_paragraph('2. Image is processed using Sharp/Imagemin libraries')
    doc.add_paragraph('3. Optimized image is stored in AWS S3')
    doc.add_paragraph('4. Download link is generated and returned to user')
    doc.add_paragraph('5. Usage statistics are updated in database')
    doc.add_paragraph('6. Redis cache stores frequently accessed data')
    
    doc.add_heading('Security Architecture', level=2)
    security_features = [
        'JWT-based authentication',
        'API key management',
        'Rate limiting and DDoS protection',
        'Input validation and sanitization',
        'HTTPS enforcement',
        'Secure file upload handling',
        'Environment variable protection'
    ]
    
    for feature in security_features:
        doc.add_paragraph(feature, style='List Bullet')

def add_technology_stack(doc):
    """Add technology stack section"""
    doc.add_heading('4. Technology Stack', level=1)
    
    doc.add_heading('Backend Technologies', level=2)
    backend_tech = [
        'Node.js 18+ - JavaScript runtime',
        'Express.js - Web framework',
        'MongoDB - NoSQL database',
        'Mongoose - MongoDB ODM',
        'Redis - In-memory cache',
        'Sharp - High-performance image processing',
        'Imagemin - Image optimization',
        'Multer - File upload handling',
        'JWT - Authentication',
        'Stripe - Payment processing',
        'AWS SDK - Cloud services integration',
        'Winston - Logging',
        'Nodemailer - Email service'
    ]
    
    for tech in backend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_heading('Frontend Technologies', level=2)
    frontend_tech = [
        'Next.js 14 - React framework',
        'React 18 - UI library',
        'TypeScript - Type safety',
        'Tailwind CSS - Utility-first CSS',
        'Framer Motion - Animations',
        'React Query - Data fetching',
        'Zustand - State management',
        'React Dropzone - File upload',
        'Stripe Elements - Payment forms',
        'Recharts - Data visualization'
    ]
    
    for tech in frontend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_heading('DevOps & Infrastructure', level=2)
    devops_tech = [
        'Docker - Containerization',
        'Docker Compose - Multi-container orchestration',
        'GitHub Actions - CI/CD',
        'Render - Backend hosting',
        'Vercel - Frontend hosting',
        'MongoDB Atlas - Cloud database',
        'AWS S3 - Object storage',
        'Redis Cloud - Managed Redis'
    ]
    
    for tech in devops_tech:
        doc.add_paragraph(tech, style='List Bullet')

def add_installation_setup(doc):
    """Add installation and setup section"""
    doc.add_heading('5. Installation & Setup', level=1)
    
    doc.add_heading('Prerequisites', level=2)
    prerequisites = [
        'Node.js 18+ installed',
        'MongoDB instance (local or cloud)',
        'AWS S3 bucket configured',
        'Stripe account for payments',
        'Redis instance (optional, falls back to memory)'
    ]
    
    for prereq in prerequisites:
        doc.add_paragraph(prereq, style='List Bullet')
    
    doc.add_heading('Quick Installation', level=2)
    doc.add_paragraph('1. Clone the repository:')
    code = doc.add_paragraph('git clone <repository-url>', style='CodeStyle')
    code.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph('2. Install dependencies:')
    code = doc.add_paragraph('npm install && cd client && npm install', style='CodeStyle')
    code.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph('3. Configure environment:')
    code = doc.add_paragraph('cp .env.example .env', style='CodeStyle')
    code.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_heading('Environment Configuration', level=2)
    doc.add_paragraph('Create a .env file with the following variables:')
    
    env_vars = [
        'PORT=5000 - Server port',
        'NODE_ENV=development - Environment mode',
        'MONGODB_URI=mongodb://localhost:27017/pixelsqueeze - Database connection',
        'JWT_SECRET=your-jwt-secret - JWT signing secret',
        'AWS_ACCESS_KEY_ID=your-access-key - AWS credentials',
        'AWS_SECRET_ACCESS_KEY=your-secret-key - AWS credentials',
        'AWS_REGION=us-east-1 - AWS region',
        'AWS_S3_BUCKET=your-bucket-name - S3 bucket name',
        'STRIPE_SECRET_KEY=sk_test_... - Stripe secret key',
        'STRIPE_WEBHOOK_SECRET=whsec_... - Stripe webhook secret',
        'REDIS_URL=redis://localhost:6379 - Redis connection (optional)'
    ]
    
    for var in env_vars:
        doc.add_paragraph(var, style='List Bullet')
    
    doc.add_heading('Running the Application', level=2)
    doc.add_paragraph('Development mode:')
    code = doc.add_paragraph('npm run dev', style='CodeStyle')
    code.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph('This starts both backend (port 5000) and frontend (port 3000)')
    
    doc.add_paragraph('Docker mode:')
    code = doc.add_paragraph('docker-compose up -d', style='CodeStyle')
    code.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_api_documentation(doc):
    """Add API documentation section"""
    doc.add_heading('6. API Documentation', level=1)
    
    doc.add_heading('Authentication', level=2)
    doc.add_paragraph('All API requests require an API key in the Authorization header:')
    code = doc.add_paragraph('Authorization: Bearer YOUR_API_KEY', style='CodeStyle')
    code.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_heading('Rate Limiting', level=2)
    doc.add_paragraph('API requests are rate-limited based on your subscription plan:')
    rate_limits = [
        'Free: 100 requests per 15 minutes',
        'Starter: 1,000 requests per 15 minutes',
        'Pro: 5,000 requests per 15 minutes',
        'Enterprise: 25,000 requests per 15 minutes'
    ]
    
    for limit in rate_limits:
        doc.add_paragraph(limit, style='List Bullet')
    
    doc.add_heading('Endpoints', level=2)
    
    doc.add_heading('POST /api/optimize', level=3)
    doc.add_paragraph('Upload and optimize an image file.')
    doc.add_paragraph('Request:')
    doc.add_paragraph('- Content-Type: multipart/form-data', style='List Bullet')
    doc.add_paragraph('- Body: image file', style='List Bullet')
    doc.add_paragraph('- Optional: quality parameter (1-100)', style='List Bullet')
    
    doc.add_heading('POST /api/optimize-url', level=3)
    doc.add_paragraph('Optimize an image from a URL.')
    doc.add_paragraph('Request body:')
    code = doc.add_paragraph('{"imageUrl": "https://example.com/image.jpg", "quality": 80}', style='CodeStyle')
    code.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_heading('GET /api/stats', level=3)
    doc.add_paragraph('Get user\'s monthly usage statistics.')
    
    doc.add_heading('Response Format', level=2)
    doc.add_paragraph('All API responses follow this format:')
    code = doc.add_paragraph('''{
  "success": true,
  "data": {
    "originalSize": 1024000,
    "optimizedSize": 256000,
    "compressionRatio": 75,
    "downloadUrl": "https://s3.amazonaws.com/...",
    "expiresAt": "2024-01-01T12:00:00Z"
  }
}''', style='CodeStyle')
    code.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_frontend_components(doc):
    """Add frontend components section"""
    doc.add_heading('7. Frontend Components', level=1)
    
    doc.add_heading('Core Components', level=2)
    components = [
        'Layout.tsx - Main application layout with navigation',
        'MarketingLayout.tsx - Marketing page layout',
        'ImageUploader.tsx - Drag-and-drop image upload interface',
        'AdvancedImageUploader.tsx - Enhanced upload with options',
        'BeforeAfter.tsx - Image comparison component',
        'RecentImages.tsx - User\'s image history',
        'StatsCard.tsx - Usage statistics display',
        'BillingPlans.tsx - Subscription plan selection',
        'Button.tsx - Reusable button component',
        'AuthProvider.tsx - Authentication context provider'
    ]
    
    for component in components:
        doc.add_paragraph(component, style='List Bullet')
    
    doc.add_heading('Page Components', level=2)
    pages = [
        'index.tsx - Landing page with hero section',
        'dashboard.tsx - User dashboard with statistics',
        'images.tsx - Image management page',
        'advanced-tools.tsx - Advanced optimization tools',
        'billing.tsx - Subscription and billing management',
        'settings.tsx - User account settings',
        'login.tsx - User authentication',
        'register.tsx - User registration',
        'features.tsx - Platform features showcase',
        'pricing.tsx - Pricing plans display'
    ]
    
    for page in pages:
        doc.add_paragraph(page, style='List Bullet')
    
    doc.add_heading('State Management', level=2)
    doc.add_paragraph('The application uses Zustand for state management with the following stores:')
    stores = [
        'authStore.ts - Authentication and user state',
        'Global state for user preferences',
        'Local state for component-specific data'
    ]
    
    for store in stores:
        doc.add_paragraph(store, style='List Bullet')

def add_backend_services(doc):
    """Add backend services section"""
    doc.add_heading('8. Backend Services', level=1)
    
    doc.add_heading('Core Services', level=2)
    services = [
        'imageProcessor.js - Core image optimization logic',
        'advancedImageProcessor.js - Advanced optimization features',
        'storageService.js - AWS S3 file management',
        'stripeService.js - Payment processing and webhooks',
        'emailService.js - Transactional email sending',
        'cleanupJob.js - Automated cleanup of expired files',
        'sentry.js - Error tracking and monitoring',
        'logger.js - Centralized logging service'
    ]
    
    for service in services:
        doc.add_paragraph(service, style='List Bullet')
    
    doc.add_heading('API Routes', level=2)
    routes = [
        'auth.js - User authentication endpoints',
        'api.js - Core image optimization API',
        'advancedImage.js - Advanced optimization features',
        'billing.js - Subscription management',
        'stripe.js - Stripe webhook handling',
        'admin.js - Administrative functions',
        'developer.js - Developer API access',
        'webhooks.js - External service webhooks'
    ]
    
    for route in routes:
        doc.add_paragraph(route, style='List Bullet')
    
    doc.add_heading('Middleware', level=2)
    middleware = [
        'auth.js - JWT authentication verification',
        'errorHandler.js - Global error handling',
        'Rate limiting middleware',
        'CORS configuration',
        'Request validation'
    ]
    
    for mw in middleware:
        doc.add_paragraph(mw, style='List Bullet')

def add_database_models(doc):
    """Add database models section"""
    doc.add_heading('9. Database Models', level=1)
    
    doc.add_heading('User Model', level=2)
    doc.add_paragraph('Stores user account information and subscription details:')
    user_fields = [
        'email - User email address (unique)',
        'password - Hashed password',
        'firstName, lastName - User name',
        'subscription - Current subscription plan',
        'apiKey - API access key',
        'usage - Monthly usage statistics',
        'createdAt, updatedAt - Timestamps'
    ]
    
    for field in user_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    doc.add_heading('Image Model', level=2)
    doc.add_paragraph('Tracks processed images and optimization results:')
    image_fields = [
        'userId - Reference to user',
        'originalName - Original filename',
        'originalSize - Original file size',
        'optimizedSize - Optimized file size',
        'compressionRatio - Compression percentage',
        'downloadUrl - S3 download link',
        'expiresAt - Link expiration time',
        'status - Processing status',
        'createdAt - Processing timestamp'
    ]
    
    for field in image_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    doc.add_heading('Database Indexes', level=2)
    indexes = [
        'email (unique) - Fast user lookup',
        'apiKey (unique) - API authentication',
        'userId + createdAt - User image history',
        'expiresAt - Cleanup job optimization'
    ]
    
    for index in indexes:
        doc.add_paragraph(index, style='List Bullet')

def add_payment_integration(doc):
    """Add payment integration section"""
    doc.add_heading('10. Payment Integration', level=1)
    
    doc.add_heading('Stripe Integration', level=2)
    doc.add_paragraph('PixelSqueeze uses Stripe for subscription management and payment processing.')
    
    doc.add_heading('Subscription Plans', level=2)
    plans = [
        'Free Plan: No charge, limited usage',
        'Starter Plan: $9/month for 5,000 images',
        'Pro Plan: $29/month for 20,000 images',
        'Enterprise Plan: $99/month for 100,000 images'
    ]
    
    for plan in plans:
        doc.add_paragraph(plan, style='List Bullet')
    
    doc.add_heading('Webhook Events', level=2)
    webhooks = [
        'customer.subscription.created - New subscription',
        'customer.subscription.updated - Plan changes',
        'customer.subscription.deleted - Cancellation',
        'invoice.payment_succeeded - Successful payment',
        'invoice.payment_failed - Failed payment'
    ]
    
    for webhook in webhooks:
        doc.add_paragraph(webhook, style='List Bullet')
    
    doc.add_heading('Payment Flow', level=2)
    doc.add_paragraph('1. User selects subscription plan')
    doc.add_paragraph('2. Stripe Elements form collects payment details')
    doc.add_paragraph('3. Payment is processed through Stripe')
    doc.add_paragraph('4. Webhook updates user subscription in database')
    doc.add_paragraph('5. User receives confirmation email')
    doc.add_paragraph('6. Usage limits are updated based on plan')

def add_deployment(doc):
    """Add deployment section"""
    doc.add_heading('11. Deployment', level=1)
    
    doc.add_heading('Production Deployment', level=2)
    doc.add_paragraph('The application is deployed using a multi-platform approach:')
    
    deployment_info = [
        'Backend: Deployed on Render.com',
        'Frontend: Deployed on Vercel',
        'Database: MongoDB Atlas (cloud)',
        'Storage: AWS S3 for file storage',
        'Cache: Redis Cloud (optional)',
        'CDN: Vercel Edge Network for frontend'
    ]
    
    for info in deployment_info:
        doc.add_paragraph(info, style='List Bullet')
    
    doc.add_heading('Docker Deployment', level=2)
    doc.add_paragraph('Local and development deployment using Docker:')
    
    docker_commands = [
        'Build: docker build -t pixelsqueeze .',
        'Run: docker run -p 3000:3000 -p 5000:5000 pixelsqueeze',
        'Compose: docker-compose up -d',
        'Stop: docker-compose down'
    ]
    
    for cmd in docker_commands:
        doc.add_paragraph(cmd, style='CodeStyle')
    
    doc.add_heading('Environment Variables', level=2)
    doc.add_paragraph('Production environment variables must be configured in deployment platforms:')
    
    prod_vars = [
        'NODE_ENV=production',
        'MONGODB_URI - Production database',
        'AWS credentials - S3 access',
        'Stripe keys - Live payment processing',
        'CORS_ORIGIN - Allowed frontend domains',
        'SENTRY_DSN - Error tracking'
    ]
    
    for var in prod_vars:
        doc.add_paragraph(var, style='List Bullet')

def add_testing(doc):
    """Add testing section"""
    doc.add_heading('12. Testing', level=1)
    
    doc.add_heading('Testing Strategy', level=2)
    doc.add_paragraph('PixelSqueeze implements a comprehensive testing approach:')
    
    testing_types = [
        'Unit Tests: Individual component testing',
        'Integration Tests: API endpoint testing',
        'End-to-End Tests: Full user workflow testing',
        'Performance Tests: Load and stress testing',
        'Security Tests: Vulnerability assessment'
    ]
    
    for test_type in testing_types:
        doc.add_paragraph(test_type, style='List Bullet')
    
    doc.add_heading('Test Commands', level=2)
    test_commands = [
        'npm test - Run all tests',
        'npm test -- --testNamePattern="API" - Run specific tests',
        'npm run test:coverage - Generate coverage report',
        'npm run test:watch - Watch mode for development'
    ]
    
    for cmd in test_commands:
        doc.add_paragraph(cmd, style='CodeStyle')
    
    doc.add_heading('Testing Tools', level=2)
    tools = [
        'Jest - JavaScript testing framework',
        'Supertest - HTTP assertion library',
        'React Testing Library - React component testing',
        'MSW - API mocking for tests'
    ]
    
    for tool in tools:
        doc.add_paragraph(tool, style='List Bullet')

def add_security(doc):
    """Add security section"""
    doc.add_heading('13. Security', level=1)
    
    doc.add_heading('Security Measures', level=2)
    security_measures = [
        'JWT Authentication - Secure token-based auth',
        'API Key Management - Unique keys per user',
        'Rate Limiting - DDoS protection',
        'Input Validation - XSS and injection prevention',
        'File Upload Security - Malware scanning',
        'HTTPS Enforcement - Encrypted communication',
        'Environment Variables - Secure configuration',
        'CORS Configuration - Cross-origin protection'
    ]
    
    for measure in security_measures:
        doc.add_paragraph(measure, style='List Bullet')
    
    doc.add_heading('Data Protection', level=2)
    data_protection = [
        'Password Hashing - bcrypt with salt',
        'API Key Encryption - Secure storage',
        'File Access Control - User-specific permissions',
        'Data Encryption - At rest and in transit',
        'Regular Security Audits - Vulnerability assessment'
    ]
    
    for protection in data_protection:
        doc.add_paragraph(protection, style='List Bullet')
    
    doc.add_heading('Compliance', level=2)
    compliance_items = [
        'GDPR Compliance - Data privacy',
        'PCI DSS - Payment security',
        'SOC 2 - Security controls',
        'Regular Penetration Testing'
    ]
    
    for item in compliance_items:
        doc.add_paragraph(item, style='List Bullet')

def add_monitoring_analytics(doc):
    """Add monitoring and analytics section"""
    doc.add_heading('14. Monitoring & Analytics', level=1)
    
    doc.add_heading('Application Monitoring', level=2)
    monitoring_features = [
        'Winston Logging - Structured logging system',
        'Sentry Integration - Error tracking and alerting',
        'Performance Metrics - Response time monitoring',
        'Health Checks - Service availability monitoring',
        'Uptime Monitoring - Service reliability tracking'
    ]
    
    for feature in monitoring_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('User Analytics', level=2)
    analytics_features = [
        'Usage Statistics - Image processing metrics',
        'User Behavior - Feature usage patterns',
        'Conversion Tracking - Subscription analytics',
        'Performance Metrics - Optimization effectiveness',
        'Error Tracking - User experience monitoring'
    ]
    
    for feature in analytics_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('Business Intelligence', level=2)
    bi_features = [
        'Revenue Analytics - Subscription metrics',
        'User Growth - Acquisition and retention',
        'Feature Usage - Popular optimization options',
        'Geographic Distribution - User location data',
        'Performance Trends - System optimization'
    ]
    
    for feature in bi_features:
        doc.add_paragraph(feature, style='List Bullet')

def add_troubleshooting(doc):
    """Add troubleshooting section"""
    doc.add_heading('15. Troubleshooting', level=1)
    
    doc.add_heading('Common Issues', level=2)
    
    doc.add_heading('Redis Connection Issues', level=3)
    doc.add_paragraph('If you see Redis connection warnings:')
    redis_solutions = [
        'Check if Redis service is running',
        'Verify Redis connection URL in environment',
        'Ensure Redis port is accessible',
        'Application will fall back to memory cache'
    ]
    
    for solution in redis_solutions:
        doc.add_paragraph(solution, style='List Bullet')
    
    doc.add_heading('Image Processing Failures', level=3)
    doc.add_paragraph('If image optimization fails:')
    image_solutions = [
        'Check file format support (JPG, PNG)',
        'Verify file size limits',
        'Ensure sufficient server memory',
        'Check Sharp/Imagemin installation'
    ]
    
    for solution in image_solutions:
        doc.add_paragraph(solution, style='List Bullet')
    
    doc.add_heading('Payment Processing Issues', level=3)
    doc.add_paragraph('If Stripe payments fail:')
    payment_solutions = [
        'Verify Stripe API keys are correct',
        'Check webhook endpoint configuration',
        'Ensure webhook secret is properly set',
        'Verify subscription plan configuration'
    ]
    
    for solution in payment_solutions:
        doc.add_paragraph(solution, style='List Bullet')
    
    doc.add_heading('Debug Mode', level=2)
    doc.add_paragraph('Enable debug logging for troubleshooting:')
    code = doc.add_paragraph('NODE_ENV=development DEBUG=* npm run dev', style='CodeStyle')
    code.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_development_guidelines(doc):
    """Add development guidelines section"""
    doc.add_heading('16. Development Guidelines', level=1)
    
    doc.add_heading('Code Standards', level=2)
    standards = [
        'Use TypeScript for type safety',
        'Follow ESLint configuration',
        'Write meaningful commit messages',
        'Include JSDoc comments for functions',
        'Use consistent naming conventions',
        'Implement proper error handling'
    ]
    
    for standard in standards:
        doc.add_paragraph(standard, style='List Bullet')
    
    doc.add_heading('Git Workflow', level=2)
    workflow = [
        'Create feature branches from main',
        'Use descriptive branch names',
        'Write clear pull request descriptions',
        'Ensure all tests pass before merging',
        'Squash commits for clean history'
    ]
    
    for step in workflow:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_heading('Testing Requirements', level=2)
    testing_reqs = [
        'Write tests for new features',
        'Maintain minimum 80% code coverage',
        'Test both success and error scenarios',
        'Include integration tests for APIs',
        'Test edge cases and boundary conditions'
    ]
    
    for req in testing_reqs:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading('Performance Guidelines', level=2)
    performance_guidelines = [
        'Optimize database queries',
        'Implement proper caching strategies',
        'Use async/await for I/O operations',
        'Monitor memory usage and leaks',
        'Implement rate limiting for APIs'
    ]
    
    for guideline in performance_guidelines:
        doc.add_paragraph(guideline, style='List Bullet')

if __name__ == "__main__":
    create_documentation()
